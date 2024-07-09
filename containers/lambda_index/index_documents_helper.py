# Copyright Amazon.com, Inc. or its affiliates. All Rights Reserved.
# SPDX-License-Identifier: MIT-0

import boto3
import json
import os
from io import BytesIO
from pypdf import PdfReader
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from opensearchpy import OpenSearch, RequestsHttpConnection, AWSV4SignerAuth, helpers
from opensearchpy.helpers import bulk

def get_s3_key_list(bucket_name, s3_prefix, file_extensions, max_file_size):
    s3_resource = boto3.resource('s3')
    key_list = []

    for key in s3_resource.Bucket(bucket_name).objects.filter(Prefix=s3_prefix):
        if key.key.endswith(file_extensions) and int(key.size) <= max_file_size:
            key_list.append(key.key)
            
    return key_list

def read_md(bucket_name, key):
    md_full_text = ""
    
    s3 = boto3.client("s3")
    md_file = s3.get_object(Bucket=bucket_name, Key=key)[
        "Body"
    ].read()
        
    return md_file.decode('utf-8')

def read_pdf(bucket_name, key):
    s3 = boto3.client("s3")
    pdf_file = s3.get_object(Bucket=bucket_name, Key=key)[
        "Body"
    ].read()
    reader = PdfReader(BytesIO(pdf_file))
    pages = reader.pages
    return pages

def read_docx(bucket_name, key):
    docx_full_text = ""

    s3 = boto3.resource('s3')
    bucket = s3.Bucket(bucket_name)
    object = bucket.Object(key)

    file_stream = BytesIO()
    object.download_fileobj(file_stream)

    document = docx.Document(file_stream)
    for para in document.paragraphs:
        docx_full_text += para.text + "\n"
        
    return docx_full_text

#Function to split and summarize a text string using a Langchain text splitter object and Titan Text Express on Bedrock
def split_and_summarize_text_until_sized(region_name, text, max_summary_length):
    text_splitter_object = RecursiveCharacterTextSplitter(
        chunk_size=15000,
        chunk_overlap=100,
        length_function=len,
    )

    bedrock_runtime_object = boto3.client(
        service_name='bedrock-runtime',
        region_name=region_name, 
    )

    llm_prompt_template = '''The following is a document:
        {text_to_summarize}
        Summarize the key points of the document in no more than 4 sentences.'''
    
    text_gen_config = {
        "maxTokenCount": 500,
        "stopSequences": [], 
        "temperature": 0,
        "topP": 1
    }
    
    model_id = 'amazon.titan-text-express-v1'
    accept = 'application/json' 
    content_type = 'application/json'
    
    while len(text) > max_summary_length:
        #print("A round of summarization. Text length is", len(text))
        new_text = ""
        sections = text_splitter_object.split_text(text)
        for section in sections:
            prompt_data = llm_prompt_template.replace("{text_to_summarize}", section)
            body = json.dumps({
            "inputText": prompt_data,
            "textGenerationConfig": text_gen_config  
            })
            response = bedrock_runtime_object.invoke_model(
                body=body, 
                modelId=model_id, 
                accept=accept, 
                contentType=content_type
            )
            response_body = json.loads(response['body'].read())
            if not "Sorry - this model is unable to" in response_body['results'][0]['outputText']:
                new_text += response_body['results'][0]['outputText'] + " "
            else:
                print(response_body['results'][0]['outputText'])
                
        text = new_text
    return text

def summarize_md_document(region_name, bucket_name, key, max_summary_length):
    text = read_md(
        bucket_name = bucket_name,
        key = key
    )
    summary_text = split_and_summarize_text_until_sized(
        region_name = region_name,
        text = text, 
        max_summary_length = max_summary_length
    )
    return summary_text

def summarize_pdf_document(region_name, bucket_name, key, max_summary_length):
    pages = read_pdf(
        bucket_name = bucket_name,
        key = key
    )
    text = ""
    for page in pages:
        text += page.extract_text()
    summary_text = split_and_summarize_text_until_sized(
        region_name = region_name,
        text = text,
        max_summary_length = max_summary_length
    )
    return summary_text

def summarize_docx_document(region_name, bucket_name, key, max_summary_length):
    text = read_docx(
        bucket_name = bucket_name,
        key = key
    )
    summary_text = split_and_summarize_text_until_sized(
        region_name = region_name,
        text = text, 
        max_summary_length = max_summary_length
    )
    return summary_text

# Function to return a list of dictionaries as OpenSearch payload given an list of S3 keys, bucket name, region, and maxiumum summary length
def summarize_documents(region_name, bucket_name, key_list, max_summary_length):

    text_splitter_object = RecursiveCharacterTextSplitter(
        chunk_size=512,
        chunk_overlap=0,
        length_function=len,
    )
    
    opensearch_payload = []

    for key in key_list:
        filename, file_extension = os.path.splitext(key)
        if file_extension == ".md":
            summary = summarize_md_document(
                region_name = region_name,
                bucket_name = bucket_name,
                key = key,
                max_summary_length = max_summary_length
            )
        elif file_extension == ".pdf":
            summary = summarize_pdf_document(
                region_name = region_name,
                bucket_name = bucket_name,
                key = key,
                max_summary_length = max_summary_length
            )
        elif file_extension == ".docx":
            summary = summarize_docx_document(
                region_name = region_name,
                bucket_name = bucket_name,
                key = key,
                max_summary_length = max_summary_length
            )

        sections = text_splitter_object.split_text(summary)
        for section_number, section in enumerate(sections):
            clean_section = section.replace(" \n", " ").replace("\n", " ")
            body = {
               "document": key,
                "section": section_number,
                "text": clean_section
            }
            
            opensearch_payload.append(body)
    return opensearch_payload

# Function to write to opensearch summary index a list of dictionaries as OpenSearch payload
def index_opensearch_summary_payload(region_name, opensearch_host, opensearch_payload, summary_index_name):
    # Get OpenSearch client
    credentials = boto3.Session().get_credentials()
    auth = AWSV4SignerAuth(credentials, region_name)
    opensearch_client = OpenSearch(
        hosts = [{'host': opensearch_host, 'port': 443}],
        http_auth = auth,
        use_ssl = True,
        verify_certs = True,
        connection_class = RequestsHttpConnection
    )

    # Define the dictionary to summarize result
    result = {
        'success_record_count': 0,
        'error_record_count': 0
        }

    # Iterate through each record and index in OpenSearch
    for item in opensearch_payload:
        result['success_record_count'] += 1
        response = opensearch_client.index(index=summary_index_name, body=item)
        if response['result'] != "created":
            print(response)

    return result

# Function to create opensearch insert dictionary from list of string from pages
def pages_to_opensearch(pages, key, opensearch_client, full_text_index_name):

    # Create a langchain text splitter object for pdf
    pdf_text_splitter_object = RecursiveCharacterTextSplitter(
        chunk_size=512,
        chunk_overlap=0,
        length_function=len,
    )

    section_number = 0
    success_record_count = 0
    error_record_count = 0
    filename, file_extension = os.path.splitext(key)
    
    for page_number, page in enumerate(pages):
        #print("Processing page", page_number)
        page_text = page.extract_text()
        sections = pdf_text_splitter_object.split_text(page_text)
        
        for section in sections:
            #print("Processing section", section_number)
            clean_section = section.replace(" \n", " ").replace("\n", " ")
            body = {
               "document": key,
                "page": page_number+1,
                "section": section_number+1,
                "text": clean_section
            }

            response = opensearch_client.index(index=full_text_index_name, body=body)
            if response['result'] == "created":
                #print("Success indexing.")
                success_record_count += 1
            else:
                #print(response)
                error_record_count += 1
            section_number += 1
    result = {
        'sections': section_number,
        'success_record_count': success_record_count,
        'error_record_count': error_record_count
    }
    return result

# Function to create opensearch insert dictionary from single text string
def text_string_to_opensearch(text, key, opensearch_client, full_text_index_name):

    success_record_count = 0
    error_record_count = 0
    filename, file_extension = os.path.splitext(key)

    if file_extension == ".md":
        # Create a langchain text splitter object for markdown and split into sections
        markdown_text_splitter_object = RecursiveCharacterTextSplitter(
            chunk_size=512,
            chunk_overlap=0,
            length_function=len,
            separators=["#"]
        )
        sections = markdown_text_splitter_object.split_text(text)

    elif file_extension == ".docx":
        # Create a langchain text splitter object for plaintext and split into sections
        plaintext_text_splitter_object = RecursiveCharacterTextSplitter(
            chunk_size=512,
            chunk_overlap=0,
            length_function=len
        )
        sections = plaintext_text_splitter_object.split_text(text)
    
    for section_number, section in enumerate(sections):
        clean_section = section.rstrip()

        section_heading = ""
        # For markdown files, try to get a section heading
        if file_extension == ".md":
            for line in section.splitlines():
                if len(line) > 2:
                    if line[0] == "#":
                        section_heading = line.strip("#").lstrip()
                        break

        body = {
            "document": key,
            "section": section_number + 1,
            "text": clean_section,
            "section_heading": section_heading
        }

        response = opensearch_client.index(index=full_text_index_name, body=body)
        if response['result'] == "created":
            success_record_count += 1
        else:
            error_record_count += 1
    
    result = {
        'sections': section_number + 1,
        'success_record_count': success_record_count,
        'error_record_count': error_record_count
    }
    return result

# Function to split and index full text from list of S3 markdown, pdf or docx keys
def split_and_index_full_text(region_name, opensearch_host, bucket_name, key_list, full_text_index_name):

    # Get OpenSearch client
    credentials = boto3.Session().get_credentials()
    auth = AWSV4SignerAuth(credentials, region_name)
    opensearch_client = OpenSearch(
        hosts = [{'host': opensearch_host, 'port': 443}],
        http_auth = auth,
        use_ssl = True,
        verify_certs = True,
        connection_class = RequestsHttpConnection
    )

    result_summary = []
    
    # Iterate through the documents in the S3 key list, read the contents and split into sections
    for count, key in enumerate(key_list):
        filename, file_extension = os.path.splitext(key)
        # Read and split markdown file
        if file_extension == ".md":
            text = read_md(bucket_name, key)
            result = text_string_to_opensearch(
                text = text,
                key = key,
                opensearch_client = opensearch_client,
                full_text_index_name = full_text_index_name
            )
        # Read and split pdf file
        elif file_extension == ".pdf":
            pages = read_pdf(bucket_name, key)
            result = pages_to_opensearch(
                pages = pages,
                key = key,
                opensearch_client = opensearch_client,
                full_text_index_name = full_text_index_name
            )
        # Read and split docx file
        elif file_extension == ".docx":
            text = read_docx(bucket_name, key)
            result = text_string_to_opensearch(
                text = text,
                key = key,
                opensearch_client = opensearch_client,
                full_text_index_name = full_text_index_name
            )
        # Not a supported file type, skip
        else:
            print(key, "- not a supported file type.", "File extension:", file_extension)
            continue

        # Append the result of this file to the result list
        result_summary.append(
            {
                'key': key,
                'sections': result['sections'],
                'success_record_count': result['success_record_count'],
                'error_record_count': result['error_record_count']
            }
        )
    return result_summary

# Function to determine a date for each file in a list and add it to the OpenSearch date index
def index_date(region_name, opensearch_host, bucket_name, key_list, date_index_name):

    documents_indexed = 0
    success_record_count = 0
    error_record_count = 0
    
    # Get OpenSearch client
    credentials = boto3.Session().get_credentials()
    auth = AWSV4SignerAuth(credentials, region_name)
    opensearch_client = OpenSearch(
        hosts = [{'host': opensearch_host, 'port': 443}],
        http_auth = auth,
        use_ssl = True,
        verify_certs = True,
        connection_class = RequestsHttpConnection
    )

    # Get S3 client and resource
    s3_client = boto3.client("s3")
    s3_resource = boto3.resource('s3')
    
    # Iterate through the documents in the S3 key list, determine a date for each and write to OpenSearch index
    # There are different ways of determining dates for files.  Update code below to suit your use case.
    for count, key in enumerate(key_list):
        filename, file_extension = os.path.splitext(key)

        # If a markdown file, date is based on S3 last modified date
        if file_extension == ".md":
            document_date = s3_client.get_object(Bucket=bucket_name, Key=key)[
                "LastModified"]            

        # If a pdf file, date is based on pdf metadata creation date
        elif file_extension == ".pdf":
            pdf_file = s3_client.get_object(Bucket=bucket_name, Key=key)[
                    "Body"
                ].read()
            reader = PdfReader(BytesIO(pdf_file))
            document_date = reader.metadata.creation_date
            
        # If a docx file, date is based on docx metadata creation date
        elif file_extension == ".docx":
            bucket = s3_resource.Bucket(bucket_name)
            object = bucket.Object(key)
            file_stream = BytesIO()
            object.download_fileobj(file_stream)
            document = docx.Document(file_stream)
            prop = document.core_properties
            document_date = prop.created

        # Not a supported file type, skip
        else:
            print(key, "- not a supported file type.", "File extension:", file_extension)
            continue

        documents_indexed += 1
            
        # Write date to OpenSearch date index
        item = {
            "document": key,
            "document_date": document_date
        }
        response = opensearch_client.index(index=date_index_name, body=item)
        
        if response['result'] == "created":
            success_record_count += 1
        else:
            error_record_count += 1
    
    result_summary = {
        'documents': documents_indexed,
        'success_record_count': success_record_count,
        'error_record_count': error_record_count
    }
    return result_summary

# Function to delete all records for a list of document keys from OpenSearch index
def delete_index_recs_by_key_list(region_name, opensearch_host, key_list, index_name):

    # Get OpenSearch client
    credentials = boto3.Session().get_credentials()
    auth = AWSV4SignerAuth(credentials, region_name)
    opensearch_client = OpenSearch(
        hosts = [{'host': opensearch_host, 'port': 443}],
        http_auth = auth,
        use_ssl = True,
        verify_certs = True,
        connection_class = RequestsHttpConnection
    )

    results = []
    
    # Iterate through the documents in the key list, delete any records in OpenSearch index
    for key in key_list:
        response = opensearch_client.delete_by_query(index=index_name, body={
          'query': {'match_phrase':{'document': key}}
        })
        results.append(response)
        
    return results